import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
import os

def create_dxc_document():
    doc = docx.Document()

    # Define Colors
    HEX_DXC_PURPLE = "603494"
    HEX_DXC_DARK   = "3A1858"
    HEX_GREY_BG    = "F5EEFD"
    HEX_GREY_LIGHT = "F2F2F2"
    HEX_TEXT_DARK  = "333333"

    COLOR_DXC_PURPLE = RGBColor(96, 52, 148)
    COLOR_DXC_DARK   = RGBColor(58, 24, 88)
    COLOR_TEXT_DARK  = RGBColor(51, 51, 51)
    COLOR_GREY_MUTED = RGBColor(102, 102, 102)

    # Page Setup
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Base Styles Configuration
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = COLOR_TEXT_DARK

    def set_cell_background(cell, hex_color):
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
        cell._tc.get_or_add_tcPr().append(shading_elm)

    def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
        tcPr.append(tcMar)

    def add_heading_1(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(8)
        h.paragraph_format.keep_with_next = True
        run = h.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = COLOR_DXC_PURPLE
        return h

    def add_heading_2(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after = Pt(6)
        h.paragraph_format.keep_with_next = True
        run = h.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = COLOR_DXC_DARK
        return h

    def add_heading_3(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(10)
        h.paragraph_format.space_after = Pt(4)
        h.paragraph_format.keep_with_next = True
        run = h.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = COLOR_TEXT_DARK
        return h

    def add_bullet(bold_prefix, text):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(4)
        r_bold = p.add_run(bold_prefix)
        r_bold.font.bold = True
        r_bold.font.color.rgb = COLOR_DXC_PURPLE
        r_text = p.add_run(text)
        return p

    def add_callout(title, text):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.cell(0, 0)
        cell.width = Inches(6.5)
        set_cell_background(cell, HEX_GREY_BG)
        set_cell_margins(cell, top=140, bottom=140, left=200, right=200)

        # Set left border color
        tcPr = cell._tc.get_or_add_tcPr()
        borders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:left w:val="single" w:sz="36" w:space="0" w:color="{HEX_DXC_PURPLE}"/><w:top w:val="none"/><w:right w:val="none"/><w:bottom w:val="none"/></w:tcBorders>')
        tcPr.append(borders)

        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(4)
        r_title = p.add_run(f"📌 {title}\n")
        r_title.font.bold = True
        r_title.font.size = Pt(11)
        r_title.font.color.rgb = COLOR_DXC_PURPLE

        r_body = p.add_run(text)
        r_body.font.size = Pt(10.5)
        doc.add_paragraph() # spacing

    # -------------------------------------------------------------
    # TITLE PAGE
    # -------------------------------------------------------------
    p_title_space = doc.add_paragraph()
    p_title_space.paragraph_format.space_before = Pt(40)

    p_org = doc.add_paragraph()
    r_org = p_org.add_run("DXC TECHNOLOGY • CYBERSECURITY CENTER OF EXCELLENCE")
    r_org.font.size = Pt(11)
    r_org.font.bold = True
    r_org.font.color.rgb = COLOR_GREY_MUTED
    p_org.paragraph_format.space_after = Pt(12)

    p_title = doc.add_paragraph()
    r_title = p_title.add_run("DXC SOC VULN CORRELATION")
    r_title.font.size = Pt(32)
    r_title.font.bold = True
    r_title.font.color.rgb = COLOR_DXC_PURPLE
    p_title.paragraph_format.space_after = Pt(4)

    p_sub = doc.add_paragraph()
    r_sub = p_sub.add_run("Technical Architecture, Threat Intelligence Integration, & Algorithmic Correlation Documentation")
    r_sub.font.size = Pt(14)
    r_sub.font.color.rgb = COLOR_DXC_DARK
    p_sub.paragraph_format.space_after = Pt(40)

    # Decorative Line
    p_line = doc.add_paragraph()
    p_line_border = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="24" w:space="1" w:color="{HEX_DXC_PURPLE}"/></w:pBdr>')
    p_line._p.get_or_add_pPr().append(p_line_border)
    p_line.paragraph_format.space_after = Pt(40)

    # Metadata Block
    p_meta = doc.add_paragraph()
    p_meta.paragraph_format.line_spacing = 1.3
    
    runs_meta = [
        ("Document Version: ", True), ("1.0.0 (Release)\n", False),
        ("Target System: ", True), ("SOC/SIEM Threat Intelligence Correlation Web Platform\n", False),
        ("Author / Maintainer: ", True), ("DXC Cybersecurity Operations Team\n", False),
        ("Date: ", True), ("August 2026\n", False),
        ("Classification: ", True), ("Internal Technical Specification\n", False),
    ]
    for text, is_bold in runs_meta:
        r = p_meta.add_run(text)
        r.font.bold = is_bold
        if is_bold: r.font.color.rgb = COLOR_DXC_DARK

    doc.add_page_break()

    # -------------------------------------------------------------
    # SECTION 1: EXECUTIVE SUMMARY & PROJECT CONTEXT
    # -------------------------------------------------------------
    add_heading_1("1. Executive Summary & Project Context")

    p = doc.add_paragraph()
    p.add_run("Modern Enterprise Security Operations Centers (SOCs) manage thousands of detection rules across fragmented SIEMs, EDRs, and Network Detection Systems (NIDS). However, security teams consistently struggle with three operational bottlenecks:")
    
    add_bullet("1. Detection Visibility Gaps: ", "Lack of a unified view linking deployed SIEM rules to known MITRE ATT&CK Threat Actor TTPs.")
    add_bullet("2. Vulnerability Blind Spots: ", "Inability to correlate active detection signatures with recently disclosed Common Vulnerabilities and Exposures (CVEs) and public exploit vectors.")
    add_bullet("3. Heterogeneous Rule Formats: ", "Siloed rule definitions across Sigma (YAML), Elastic EQL (JSON), Splunk ES (SPL JSON), and Suricata NIDS (YAML) making comparative analysis difficult.")

    p_summary = doc.add_paragraph()
    p_summary.add_run("The ")
    r_bold = p_summary.add_run("DXC SOC VULN CORRELATION Platform")
    r_bold.font.bold = True
    r_bold.font.color.rgb = COLOR_DXC_PURPLE
    p_summary.add_run(" is a single-pane-of-glass client-side application designed to solve these challenges. It ingests detection rules in any major format, parses them into a standardized internal schema, and automatically correlates them against ")
    p_summary.add_run("MITRE ATT&CK Enterprise v15")
    p_summary.add_run(" threat intelligence and live ")
    p_summary.add_run("NVD (National Vulnerability Database) CVE data.")

    add_callout("Key Strategic Objective", 
                "Empower SOC Analysts, Threat Hunters, and Detection Engineers to instantly identify weak rules, uncover exploitable gaps targetable by specific Advanced Persistent Threat (APT) groups, and generate actionable correlation intelligence reports.")

    # -------------------------------------------------------------
    # SECTION 2: SYSTEM ARCHITECTURE & RESOURCES
    # -------------------------------------------------------------
    add_heading_1("2. System Architecture & Resources")

    p_arch = doc.add_paragraph()
    p_arch.add_run("The platform is engineered as a 100% client-side Single-Page Application (SPA), guaranteeing zero server footprint, strict data privacy (user rules never leave the browser), and high speed.")

    add_heading_2("2.1 Technology Stack")

    tbl_tech = doc.add_table(rows=5, cols=3)
    tbl_tech.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_tech.autofit = False

    headers = ["Layer", "Technology / Framework", "Purpose / Function"]
    hdr_cells = tbl_tech.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], HEX_DXC_PURPLE)
        p = hdr_cells[i].paragraphs[0]
        p.runs[0].font.bold = True
        p.runs[0].font.color.rgb = RGBColor(255, 255, 255)

    data_tech = [
        ("Frontend Core", "Vanilla HTML5 + JavaScript (ES2022 Modules)", "Modular SPA router, state manager, DOM manipulation"),
        ("Design System", "CSS3 Custom Tokens + DXC Brand Theme", "DXC Purple (#603494) color system, responsive grid"),
        ("Data Parsers", "js-yaml CDN + Native JSON Engine", "Universal parsing of Sigma, Elastic, Splunk, Suricata"),
        ("Graph Visualization", "D3.js v7 Force-Directed Network", "Interactive node correlation graph rendering"),
    ]

    for row_idx, data in enumerate(data_tech, start=1):
        row_cells = tbl_tech.rows[row_idx].cells
        bg_color = HEX_GREY_BG if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, text in enumerate(data):
            row_cells[col_idx].text = text
            set_cell_background(row_cells[col_idx], bg_color)
            set_cell_margins(row_cells[col_idx], top=80, bottom=80, left=120, right=120)

    doc.add_paragraph() # spacing

    add_heading_2("2.2 Primary Data Resources & Integration APIs")

    add_bullet("1. MITRE ATT&CK STIX Bundle (v15): ", "Fetched live from the official CTI GitHub repository (enterprise-attack.json ~8MB). The platform indexes 140+ APT intrusion sets, 900+ techniques/sub-techniques, 14 tactics, courses of action (mitigations), and malware/tool relationships into indexed in-memory lookup maps.")
    add_bullet("2. NIST NVD CVE REST API v2.0: ", "Interfaced client-side via https://services.nvd.nist.gov/rest/json/cves/2.0. Supports live keyword queries, CWE category queries, publication date ranges, and direct CVE ID lookup. Includes rate-limiting controls and a 15+ notable baseline CVE fallback for offline resilience.")
    add_bullet("3. Rule Sample Dataset: ", "Includes 4 downloadable production-grade example rule sets in Sigma YAML, Elastic EQL JSON, Splunk ES JSON, and Suricata NIDS YAML formats.")

    # -------------------------------------------------------------
    # SECTION 3: DETAILED FEATURE BREAKDOWN (THE 4 SECTIONS)
    # -------------------------------------------------------------
    add_heading_1("3. Detailed Feature Breakdown")

    add_heading_2("3.1 Section 1: CVE Intelligence Engine")
    p_cve = doc.add_paragraph()
    p_cve.add_run("The CVE Section enables security teams to research public vulnerabilities, analyze CVSS severity vectors, and inspect affected software systems.")
    add_bullet("Direct Code Name Lookup: ", "Query any vulnerability directly by ID (e.g., CVE-2021-44228, CVE-2023-34362).")
    add_bullet("Category / CWE Filter: ", "Filter vulnerabilities by weakness classification (XSS CWE-79, SQLi CWE-89, Buffer Overflow CWE-119, Deserialization CWE-502, etc.).")
    add_bullet("CVSS Score Sliders: ", "Dynamic dual-slider control filtering base scores from 0.0 to 10.0.")
    add_bullet("Date & Product Filters: ", "Filter by publication date ranges and inspect affected CPE product strings.")
    add_bullet("CVSS Metric Visualizer: ", "SVG score rings, severity badges (CRITICAL, HIGH, MEDIUM, LOW), CVSS v3.1 vector strings, and official NVD references.")

    add_heading_2("3.2 Section 2: APT Groups & TTP Browser")
    p_apt = doc.add_paragraph()
    p_apt.add_run("The APT Section provides complete visibility into threat actor groups and their tactics, techniques, and procedures (TTPs).")
    add_bullet("APT Group Search: ", "Instant searchable catalog of 140+ Advanced Persistent Threat groups by primary name, country attribution, or alias (e.g., APT29, Cozy Bear, LockBit, Cobalt Strike).")
    add_bullet("TTP Breakdown: ", "Selecting any group displays all associated techniques grouped by MITRE ATT&CK tactics (Initial Access, Execution, Persistence, Privilege Escalation, C2, etc.).")
    add_bullet("Detailed TTP Modal: ", "Clicking any technique opens a deep-dive panel containing technique descriptions, detection guidance, data sources, course-of-action mitigations, and known malware tools.")

    add_heading_2("3.3 Section 3: SOC/SIEM Rules Visualizer & Comparator")
    p_rules = doc.add_paragraph()
    p_rules.add_run("This section handles multi-format rule ingestion and comparative diff analysis.")
    add_bullet("Visualizer Mode: ", "Upload any Sigma, Elastic, Splunk, or Suricata file. The parser normalizes rules into unified visual cards highlighting detection logic, status, severity, tags, and MITRE mappings.")
    add_bullet("Comparator Mode: ", "Upload two rule files side-by-side (e.g., File A vs. File B). The comparator engine performs exact logic hashing and technique overlap analysis to display:")
    add_bullet("  • Common Rules: ", "Rules identical in both files.")
    add_bullet("  • Unique Rules (File A / B): ", "Rules present only in one file.")
    add_bullet("  • Modified Rules: ", "Rules sharing a name but exhibiting different detection logic or severity.")
    add_bullet("  • Overlap Metrics: ", "Visual percentage progress bar and overlap stats.")

    add_heading_2("3.4 Section 4: Core Rule Correlation Engine")
    p_corr = doc.add_paragraph()
    p_corr.add_run("The core objective of the application. The Correlation Engine links ingested SIEM rules directly against MITRE ATT&CK techniques, APT threat groups, and CVE vulnerability data.")
    
    # -------------------------------------------------------------
    # SECTION 4: CORRELATION METHODOLOGY & ALGORITHMIC SCORING
    # -------------------------------------------------------------
    add_heading_1("4. Correlation Methodology & Algorithmic Scoring")

    p_method = doc.add_paragraph()
    p_method.add_run("The correlation process operates across three distinct scan modes:")

    tbl_modes = doc.add_table(rows=4, cols=3)
    tbl_modes.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers_m = ["Scan Mode", "Algorithmic Scope", "Performance Profile"]
    for i, title in enumerate(headers_m):
        tbl_modes.rows[0].cells[i].text = title
        set_cell_background(tbl_modes.rows[0].cells[i], HEX_DXC_PURPLE)
        p = tbl_modes.rows[0].cells[i].paragraphs[0]
        p.runs[0].font.bold = True
        p.runs[0].font.color.rgb = RGBColor(255, 255, 255)

    data_modes = [
        ("⚡ Simple / Fast", "Direct technique ID tag matching + embedded rule CVE references.", "Fast (~5–10s)"),
        ("🔍 Advanced / Normal", "ATT&CK technique ID mapping + NVD keyword matching against rule detection logic.", "Normal (~20–30s)"),
        ("🧬 Extensive / Deep", "Full ATT&CK cross-referencing (Campaigns, Sub-techniques, APT frequency weighting, CVSS scoring).", "Deep (~1–2 min)"),
    ]

    for row_idx, data in enumerate(data_modes, start=1):
        row_cells = tbl_modes.rows[row_idx].cells
        bg_color = HEX_GREY_BG if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, text in enumerate(data):
            row_cells[col_idx].text = text
            set_cell_background(row_cells[col_idx], bg_color)
            set_cell_margins(row_cells[col_idx], top=80, bottom=80, left=120, right=120)

    doc.add_paragraph() # spacing

    add_heading_2("4.1 Mathematical Scoring Formula")
    p_math = doc.add_paragraph()
    p_math.add_run("Each rule receives an Overall Correlation Score S (0.0 to 1.0) calculated via:")
    
    add_callout("Correlation Score Formula",
                "S_overall = Min( 1.0,  S_technique + S_CVE + S_risk )\n\n"
                "• S_technique = Average confidence of matched techniques (Direct tag = 1.0, Keyword = 0.4–0.85, Tactic = 0.3)\n"
                "• S_CVE = Matched CVE Count × 0.05 (Max 0.20)\n"
                "• S_risk = Rule Risk Score / 200 (Max 0.25)")

    add_heading_2("4.2 Rule Coverage Classification")
    add_bullet("🔴 Weak / Uncovered Rules (Score < 30%): ", "Rules with no ATT&CK technique mapping or low confidence logic. High risk of security blind spot.")
    add_bullet("🟠 Partial Coverage Rules (Score 30% – 65%): ", "Rules mapped to techniques with moderate confidence or missing CVE references.")
    add_bullet("🟢 Well Correlated Rules (Score ≥ 65%): ", "High-confidence rules fully mapped to ATT&CK techniques, APT threat groups, and CVE vulnerability vectors.")

    add_heading_2("4.3 Interactive D3 Force-Directed Correlation Graph")
    p_graph = doc.add_paragraph()
    p_graph.add_run("Correlation output is rendered as a 4-tier interactive force-directed network graph:")
    add_bullet("Purple Nodes (Rules): ", "Represent ingested SIEM detection rules. Weak rules display an animated pulsing red ring.")
    add_bullet("Blue Nodes (Techniques): ", "Represent mapped MITRE ATT&CK techniques.")
    add_bullet("Red Nodes (APT Groups): ", "Represent threat actor groups known to exploit those techniques.")
    add_bullet("Orange Nodes (CVEs): ", "Represent specific vulnerabilities linked to the rule or technique.")

    # -------------------------------------------------------------
    # SECTION 5: EXAMPLE RULE FILES SECTION
    # -------------------------------------------------------------
    add_heading_1("5. Example Rule Files Section")

    p_ex = doc.add_paragraph()
    p_ex.add_run("To facilitate immediate testing without requiring existing SIEM rule exports, the webapp includes an ")
    p_ex.add_run("Example Files Section").bold = True
    p_ex.add_run(" featuring 4 downloadable sample rule sets:")

    tbl_ex = doc.add_table(rows=5, cols=3)
    tbl_ex.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers_e = ["Filename", "Format", "Included Threat Scenarios"]
    for i, title in enumerate(headers_e):
        tbl_ex.rows[0].cells[i].text = title
        set_cell_background(tbl_ex.rows[0].cells[i], HEX_DXC_PURPLE)
        p = tbl_ex.rows[0].cells[i].paragraphs[0]
        p.runs[0].font.bold = True
        p.runs[0].font.color.rgb = RGBColor(255, 255, 255)

    data_ex = [
        ("sigma_windows.yml", "Sigma YAML", "Mimikatz LSASS access, Pass-the-Hash, Encoded PowerShell, Schtasks, NTDS.dit access, WMI lateral movement."),
        ("elastic_siem.json", "Elastic EQL JSON", "Cobalt Strike beaconing, LOLBin discovery, Volume Shadow Copy ransomware deletion, DNS tunneling, ProxyShell exploitation."),
        ("splunk_spl.json", "Splunk ES JSON", "C2 periodic beaconing, Large outbound data exfiltration, Password spraying brute force, Tor exit node C2, Kerberoasting."),
        ("suricata_nids.yaml", "Suricata NIDS YAML", "Emotet C2 loader, Apache Log4Shell (CVE-2021-44228), Nmap SYN scan, APT29 WellMess RAT, SMBGhost (CVE-2020-0796)."),
    ]

    for row_idx, data in enumerate(data_ex, start=1):
        row_cells = tbl_ex.rows[row_idx].cells
        bg_color = HEX_GREY_BG if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, text in enumerate(data):
            row_cells[col_idx].text = text
            set_cell_background(row_cells[col_idx], bg_color)
            set_cell_margins(row_cells[col_idx], top=80, bottom=80, left=120, right=120)

    doc.add_paragraph() # spacing

    # -------------------------------------------------------------
    # SECTION 6: DEPLOYMENT & USER GUIDE
    # -------------------------------------------------------------
    add_heading_1("6. Deployment & User Guide")

    add_heading_2("6.1 Running the Webapp Locally")
    p_dep = doc.add_paragraph()
    p_dep.add_run("Because the platform uses ES2022 JavaScript modules, it should be served via a local web server:")
    
    p_cmd = doc.add_paragraph()
    r_code = p_cmd.add_run("python -m http.server 8765 --directory C:\\path\\to\\dxc-soc-vuln-correlation")
    r_code.font.name = 'Consolas'
    r_code.font.size = Pt(10)
    r_code.font.color.rgb = COLOR_DXC_PURPLE
    p_cmd.paragraph_format.left_indent = Inches(0.4)

    p_open = doc.add_paragraph()
    p_open.add_run("Then open your browser and navigate to ")
    p_open.add_run("http://localhost:8765").bold = True

    add_heading_2("6.2 Exporting Reports")
    p_exp = doc.add_paragraph()
    p_exp.add_run("After running any correlation scan, click ")
    p_exp.add_run("Export Correlation Report (JSON)").bold = True
    p_exp.add_run(" to download a structured report containing rule metrics, top threat actors, weak rule flags, and matched CVE lists suitable for ingestion into ticketing systems (Jira, ServiceNow).")

    # Save document
    filename = "DXC_SOC_VULN_CORRELATION_Documentation.docx"
    filepath = os.path.join(r"C:\Users\hachn\.gemini\antigravity\scratch\dxc-soc-vuln-correlation", filename)
    doc.save(filepath)
    print(f"Document successfully created at: {filepath}")
    return filepath

if __name__ == "__main__":
    create_dxc_document()
